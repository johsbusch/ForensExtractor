# This program was developed by Johannes Busch @ Dept. of Forensic Medicine, University of Copenhagen. Please Google me for contact information. This version is dated Jan 30th 2025.
# This script is a custom program that imports docx-files and mines various data from them. It was developed as a tool to perform large-scale data extraction from historic autopsy records.
# I am not in any way formally trained in software programming, this code has been written in spare time and partially using LLM-tools for assistance. I do not guarantee for the functionality or validity of data extracted using the program.
# The code is tailored to Danish records, but in principle all search terms can be modified to another language. The code is not very self-explanatory. Please contact me if you have any questions.
# Feel free to modify the code for use on your own records and contact me with ideas for collaboration. This tool was developed specifically with the aspiration of multi-center cooperation studies, enabling large and diverse autopsy data sets.

# This section imports revelant packages to Python.
import os
import re
import csv
import time
from docx import Document
from collections import OrderedDict
import fitz # PyMuPDF


def search_for_COD_keywords(doc_text, regex_dict):
    paragraphs = re.split(r"\.\s*", doc_text)
    found_keywords = {label: False for label in regex_dict}

    for i, paragraph in enumerate(paragraphs):
        if "dødsårsag" in paragraph.lower():
            context = []
            context.append(paragraph)
            if i + 1 < len(paragraphs):
                context.append(paragraphs[i + 1])
            if i + 2 < len(paragraphs):
                context.append(paragraphs[i + 2])
            context_text = " ".join(context)

            for label, pattern in regex_dict.items():
                if re.search(pattern, context_text, re.IGNORECASE):
                    found_keywords[label] = True

    return found_keywords


def store_COD_text(doc_text):
    COD_pattern = re.compile(r"\bdødsårsag\w*\b", re.IGNORECASE)
    textCOD = []
    paragraphs = re.split(r"\.\s*", doc_text)

    for i, paragraph in enumerate(paragraphs):
        if COD_pattern.search(paragraph.lower()):
            # Include the current paragraph
            textCOD.append(paragraph)
            # Include the next two paragraphs if they exist
            if i + 1 < len(paragraphs):
                textCOD.append(paragraphs[i + 1])
            if i + 2 < len(paragraphs):
                textCOD.append(paragraphs[i + 2])
            textCOD = " ".join(textCOD)
            break

    return textCOD

def store_vaccine_text(doc):
    vac_pattern = re.compile(r"\b(vaccin\w*)\b", re.IGNORECASE)

    textVAC = ""
    num = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            match = vac_pattern.search(run.text)
            if match:
                num = num + 1
                if num < 2:
                    textVAC = textVAC + run.text
                else:
                    textVAC = textVAC + " / " + run.text
                    textVAC = textVAC.replace("\n", " ").replace("\r", " ")

    return textVAC


def store_finde_text(doc):
    text_finde = ""

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.lower()
        if (
            paragraph.text.lower().startswith("af sagsakterne fremgår")
            or paragraph.text.lower().startswith("af disse papirer")
            or paragraph.text.lower().startswith("nu afdøde")
            or paragraph.text.lower().startswith("der foreligger rapport fra")
            or paragraph.text.lower().startswith("om hændelsesforløbet")
            or paragraph.text.lower().startswith("det fremgår af det foreliggende")
            or paragraph.text.lower().startswith("det fremgår")
            or paragraph.text.lower().startswith("af det foreliggende fremgår")
            ):
            if text_finde != "":
                text_finde = text_finde + " / " + paragraph.text
            else:
                text_finde = paragraph.text
            # Replace newlines and other unwanted whitespace characters with a space
            text_finde = text_finde.replace("\n", " ").replace("\r", " ")

    return text_finde
            

def findeomst(doc, regex_dict):
    found_patterns = {label: False for label in regex_dict}

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.lower()
        if (
            paragraph.text.lower().startswith("af sagsakterne fremgår")
            or paragraph.text.lower().startswith("af disse papirer")
            or paragraph.text.lower().startswith("nu afdøde")
            or paragraph.text.lower().startswith("der foreligger rapport fra")
            or paragraph.text.lower().startswith("om hændelsesforløbet")
            or paragraph.text.lower().startswith("det fremgår af det foreliggende")
            or paragraph.text.lower().startswith("det fremgår")
            or paragraph.text.lower().startswith("af det foreliggende fremgår")
            ):
            for label, pattern in regex_dict.items():
                if re.search(pattern, paragraph.text, re.IGNORECASE):
                    found_patterns[label] = True
                    
    return found_patterns

# Extract "kendte sygdomme", e.g. text in conclusion occurring between "efter det oplyste" and "mand|kvinde|pige|dreng"
def kendtMed(doc):
    pattern = re.compile(r"((efter det oplyste)(.*?)(mand|kvinde|pige|dreng))", re.IGNORECASE | re.DOTALL)
    #pattern = re.compile(r"((oplyste))", re.IGNORECASE | re.DOTALL)

    textKM = ""
    
    for paragraph in doc.paragraphs:
        match = pattern.search(paragraph.text)
        if match:
            textKM = match.group(1)
            break

    return textKM

def hjerteText(doc):
    pattern_hjertepose = re.compile(r"((.*?)(hjerteposen)(.*?))", re.IGNORECASE | re.DOTALL)
    pattern_complete = re.compile(r"((.*?)(hjerteposen)(.*?)(farven))", re.IGNORECASE | re.DOTALL)
    pattern_farven = re.compile(r"((.*?)(farven))", re.IGNORECASE | re.DOTALL)

    textHeart = ""

    for paragraph in doc.paragraphs:
        match = pattern_complete.search(paragraph.text)
        if match:
            textHeart = paragraph.text
            break

    for paragraph in doc.paragraphs:
        match = pattern_hjertepose.search(paragraph.text)
        if match:
            textHeart = paragraph.text
            for paragraph in doc.paragraphs:
                match = pattern_farven.search(paragraph.text)
                if match:
                    textHeart += paragraph.text + " "
                    break

    return textHeart

def aortaText(doc):
    pattern_aorta = re.compile(r"Legemspulsåren og")

    textAorta = ""

    for paragraph in doc.paragraphs:
        match = pattern_aorta.search(paragraph.text)
        if match:
            textAorta = paragraph.text
            break
        
    return textAorta

def carotidText(doc):
    pattern_carotid = re.compile(r"Halspulsårerne")

    textCarotid = ""

    for paragraph in doc.paragraphs:
        match = pattern_carotid.search(paragraph.text)
        if match:
            textCarotid = paragraph.text
            break
        
    return textCarotid

# Extract paragraph text where "skumsvamp" occurs - negative lookbehind removes any case, where "skumsvamp" is preceded by either "ingen" or "ikke" or "eller"
def skumsvampPara(doc):
    skum_pattern = re.compile(r"(?<!ingen)(?<!ikke)(?<!eller) \b(skumsvamp\w*)\b", re.IGNORECASE)

    skum_para = ""
    num = 0

    for paragraph in doc.paragraphs:
        match = skum_pattern.search(paragraph.text)
        if match:
            num = num + 1
            if num < 2:
                skum_para = paragraph.text
            else:
                skum_para = skum_para + " / " + paragraph.text

    return skum_para

# Extract sentence where "strip" occurs, adds each occurence to a single string, with " / " between each.
def stripPara(doc):
    strip_pattern = re.compile(r"strip", re.IGNORECASE)

    strip_text = ""

    for paragraph in doc.paragraphs:
        match = strip_pattern.search(paragraph.text)
        if match:
            strip_text = strip_text + " / " + paragraph.text

    return strip_text
    

# Extract paragraph text where "tegn på sygdom" occurs
def search_TPS(doc):
    TPS_pattern = re.compile(r"((?<=tegn på sygdom).*)", re.IGNORECASE | re.DOTALL)
    iTPS_pattern = re.compile(r"((?<=ingen tegn på sygdom).*)", re.IGNORECASE | re.DOTALL)

    TPS_para = None

    for paragraph in doc.paragraphs:
        
        match_iTPS = iTPS_pattern.search(paragraph.text)
        if match_iTPS:
            TPS_para = "ingen tegn på sygdom"
            break
        
        match = TPS_pattern.search(paragraph.text)
        if match:
            TPS_para = match.group(1)
            break

    if TPS_para is not None:    
        return TPS_para
    else:
        print ("TPS pattern not found")
        return "TPS pattern not found"

def CT_search(doc, keywordCT):

    keyCT = False

    # Loop through paragraphs for "CT"
    for i, paragraph in enumerate(doc.paragraphs):
        if re.search(r"CT", paragraph.text, re.IGNORECASE):
            # Check the next four paragraphs for keywordCT
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                if re.search(r"\b{}\w*\b".format(re.escape(keywordCT)), doc.paragraphs[j].text, re.IGNORECASE):
                    print("*" + str(keywordCT) + "*" + " found in CT paragraphs")
                    keyCT = True
                    break

    return keyCT


def extract_lung_weights(text, keywords):
    # Constructing the regular expression pattern dynamically from the list of keywords

    keyword_pattern = "|".join([re.escape(keyword) for keyword in keywords])
    pattern = re.compile(
        rf"({keyword_pattern})(?![^.]*(?:blodans|bris|væskeans|hjertepose))[^.]*?(?:ca\.)?[^.]*?(\d{{2,4}})\s*(?:gram|g)", re.IGNORECASE  #replacing "[^.]*?" with "(?:(?:[^.]*\.){0,1}[^.]*)?" allows 1 dot, but does not work right now. 
    )

    """# Search for all matches in the text, only after character number 2000, i.e. skips the conclusion and introduction (should fit almost all documents)
    matches = pattern.findall(text[2500:])"""

    # 26th August 2024 - This updated regex may allow for bypassing extra weights in the sentences, e.g. "blodansanmling der vejer 30 g" - but it does not work with findall - needs solution:
    # ({keyword_pattern})(?![^.]*(blodans|bris|væskeans))[^.]*?(\d{{2,4}}) (?:gram|g)
    # 15th August 2025 - Added [^.]*?(?:ca\.)? to the regex. This adds the non-capturing group "ca." to allow for the word "ca." to occur zero or one times. Thus "." is still not permitted, except when part of "ca."

    # Phrase to start searching from (case insensitive)
    start_phrase = "Indvendig undersøgelse"

    # Search for start phrase and store it as a start position
    start_position = text.find(start_phrase)

    try:
        if start_position != -1:
            relevant_text = text[start_position:]

            # Search for all matches in the text, only after start phrase (i.e. skips conclusion, introduction, CT, udvendig undersøgelse)
            matches = pattern.findall(relevant_text)

            # Create a dictionary to store the weights associated with the keywords
            weights = {}
            for match in matches:
                keyword, weight = match
                # Convert weight to integer
                weight = int(weight)
                keyword = keyword.lower()
                weights.setdefault(keyword, weight)

            return weights

        else:
            relevant_text = text[2500:]

            # Search for all matches in the text, only after start phrase (i.e. skips conclusion, introduction, CT, udvendig undersøgelse)
            matches = pattern.findall(relevant_text)

            # Create a dictionary to store the weights associated with the keywords
            weights = {}
            for match in matches:
                keyword, weight = match
                # Convert weight to integer
                weight = int(weight)
                keyword = keyword.lower()
                weights.setdefault(keyword, weight)

            return weights
    except Exception as e:
        print(f"Error processing weights")

def extract_organ_size(text, keywords):
    # Constructing the regular expression pattern dynamically from the list of keywords

    keyword_pattern = "|".join([re.escape(keyword) for keyword in keywords])
    pattern = re.compile(
        rf"({keyword_pattern})[^.]*?måler (\d+|\d+,\d+) x (\d+|\d+,\d+) x (\d+|\d+,\d+) cm", re.IGNORECASE  #Should match "[keyword] måler 12 x 12 x 12 cm" - commas should be accepted
    )

    """# Search for all matches in the text, only after character number 2000, i.e. skips the conclusion and introduction (should fit almost all documents)
    matches = pattern.findall(text[2500:])"""

    # Phrase to start searching from (case insensitive)
    start_phrase = "Indvendig undersøgelse"

    # Search for start phrase and store it as a start position
    start_position = text.find(start_phrase)


    try:
        if start_position != -1:
            relevant_text = text[start_position:]

            # Search for all matches in the text, only after start phrase (i.e. skips conclusion, introduction, CT, udvendig undersøgelse)
            matches = pattern.findall(relevant_text)

            # Create a dictionary to store the sizes associated with the keywords, all four groups must be matched, otherwise no match is made
            sizes = {}
            for match in matches:
                organ, height, width, depth = match
                # Convert comma-decimal values to float-compatible dot-decimals
                sizes[f"{organ}_højde"] = float(height.replace(",", "."))
                sizes[f"{organ}_bredde"] = float(width.replace(",", "."))
                sizes[f"{organ}_dybde"] = float(depth.replace(",", "."))

            return sizes

        else:
            relevant_text = text[2500:]

            # Search for all matches in the text, only after start phrase (i.e. skips conclusion, introduction, CT, udvendig undersøgelse)
            matches = pattern.findall(relevant_text)

            # Create a dictionary to store the weights associated with the keywords
            sizes = {}
            for match in matches:
                organ, height, width, depth = match
                # Convert comma-decimal values to float-compatible dot-decimals
                sizes[f"{organ}_højde"] = float(height.replace(",", "."))
                sizes[f"{organ}_bredde"] = float(width.replace(",", "."))
                sizes[f"{organ}_dybde"] = float(depth.replace(",", "."))

            return sizes
    except Exception as e:
        print(f"Error processing weights")

    
def extract_wall_thicknesses(text, keywords):
    # Constructing the regular expression pattern dynamically from the list of keywords
    keyword_pattern = "|".join([re.escape(keyword) for keyword in keywords])
    pattern = re.compile(rf"({keyword_pattern}).*?(\d{{1,2}}) mm", re.IGNORECASE)

    # Create a dictionary to store the thicknesses associated with the keywords
    thicknesses = {}

    # Search for matches in the text
    for sentence in re.split(r"\.\s*", text):
        matches = pattern.findall(sentence)
        for match in matches:
            keyword, thickness = match
            # Check if thickness is not an empty string
            if thickness:
                thicknesses.setdefault(keyword.lower(), int(thickness))

    return thicknesses

def extract_pleural_fluid(text, keywords):
    # Constructing the regular expression pattern dynamically from the list of keywords
    keyword_pattern = "|".join([re.escape(keyword) for keyword in keywords])
    pattern = re.compile(rf"(\d{{1,4}}) ml.*?({keyword_pattern})|({keyword_pattern}).*?(\d{{1,4}}) ml", re.IGNORECASE)

    # Create a dictionary to store the volumes associated with the keywords
    volumes = {keyword: None for keyword in keywords}

    # Search for matches in the text (split up sentences by "." ; "," and "og"
    for sentence in re.split(r"[.]\s*|\s+og\s+|[,]", text, flags = re.IGNORECASE):
        matches = pattern.findall(sentence)
        for match in matches:
            if match[0] and match[1]:
                # Case: "200 ml væske i højre lungehule"
                volume = int(match[0])
                keyword = match[1].lower()
            elif match[2] and match[3]:
                # Case: "I højre lungehule ses 200 ml væske"
                volume = int(match[3])
                keyword = match[2].lower()
            else:
                continue

            # Store the volume in the dictionary
            if keyword in volumes:
                volumes[keyword] = volume

    return volumes

def extract_height_weight(text):
    # Define regular expression patterns for height and weight
    height_pattern = re.compile(r"Højde(?:n)?(?: er)? (\d+)\s(cm)", re.IGNORECASE)
    weight_pattern = re.compile(r"vægt(?:en)?(?: er)?\s((?:\d+,)?\d+)\s?(kg|kilo)(\.|,|\s)", re.IGNORECASE)
    weight_pattern_g = re.compile(r"vægt(?:en)?(?: er)?.*?(\d+)\s?(g)(?:ram)?(\.|,|\s)", re.IGNORECASE)

    #height_pattern = re.compile(r"Højden er (\d+).*?cm", re.IGNORECASE)
    #weight_pattern = re.compile(r"vægten(?: er)? (\d+).*?(kg)", re.IGNORECASE)
    #weight_pattern_g = re.compile(r"vægten(?: er)? (\d+).*?(g)", re.IGNORECASE)

    # Initialize variables to store height and weight
    height = None
    weight = None
    weight_unit = None

    # Search for height in the text
    height_match = height_pattern.search(text)
    if height_match:
        height = int(height_match.group(1))

    # Search for weight in the text
    weight_match = weight_pattern.search(text)
    weight_match_g = weight_pattern_g.search(text)
    if weight_match:
        weight = weight_match.group(1)
        weight_unit = weight_match.group(2)
        #Check for cases where weight is reported as decimal number, e.g. vægt 1,234 kg - if so, replace comma with dot (1.234), change to floating point number, round this number (1.234 -> 1) and convert to integer. 
        if "," in weight:
            weight = int(round(float(weight.replace(",", "."))))
        else:
            weight = int(weight)
        if weight_unit == "kilo":
            weight_unit = "kg"
    elif weight_match_g:
        weight = int(weight_match_g.group(1))
        weight_unit = weight_match_g.group(2)

    #print(height, weight, weight_unit)
    return height, weight, weight_unit


def extract_cpr_number_from_table(doc):
    # Extract CPR number from the first table in the document
    cpr = "No CPR match"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Extract and return the CPR number (last four symbols can be any combination of digits, small letters and capital letters)
                cpr_match = re.search(r"\b\d{6}-[\da-zA-ZÅ-ø]{4}\b", cell.text)
                if cpr_match:
                    #print(cpr_match.group(0))
                    cpr = cpr_match.group(0)
                    break

    if cpr == "No CPR match":
        for paragraph in doc.paragraphs:
            old_CPR_match = re.search(r"\b(\d{6})([\da-zA-ZÅ-ø]{4})(-)?\b", paragraph.text)
            if old_CPR_match:
                #print("Old CPR match - no hyphen:" + str(old_CPR_match.group(1)) + "-" + str(old_CPR_match.group(2)))
                cpr = str(old_CPR_match.group(1)) + "-" + str(old_CPR_match.group(2))
                break
            else:
                old_CPR_match = re.search(r"\b(\d{6}-[\da-zA-ZÅ-ø]{4})(-)?\b", paragraph.text)
                if old_CPR_match:
                    #print("Old CPR match:" + str(old_CPR_match.group(0)))
                    cpr = old_CPR_match.group(0)
                    break
                else:
                    old_CPR_match = re.search(r"\b(\d{2}) ?\.?(\d{2}) ?\.?(\d{2})( ?)-( ?)(\d{4})", paragraph.text)
                    if old_CPR_match:
                        #print("Old CPR match - spaces:" + str(old_CPR_match.group(1)) + str(old_CPR_match.group(2)) + str(old_CPR_match.group(3)) + "-" + str(old_CPR_match.group(6)))
                        cpr = str(old_CPR_match.group(1)) + str(old_CPR_match.group(2)) + str(old_CPR_match.group(3)) + "-" + str(old_CPR_match.group(6))
    return cpr

def extract_aut_date_from_table(doc):
    # Extract autopsy date from the first table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Extract and return the CPR number
                date_match = re.search(r"\b((\d{2})(\-|\.)(\d{2})(\-|\.)(\d{4}))\b", cell.text)
                if date_match:
                    #print(date_match.group(0))
                    return date_match.group(0)
                    break

    for paragraph in doc.paragraphs[:15]:
        old_date_match = re.search(r"(\d{1,2})(\-|\.)(\d{2})(\-|\.)(\d{2,4})", paragraph.text)
        if old_date_match:
            old_date_dd = old_date_match.group(1)
            old_date_mm = old_date_match.group(3)
            old_date_yy = old_date_match.group(5)
            if len(old_date_dd) == 1:
                old_date_dd = str("0" + old_date_dd)
            #print(str("date match:" + old_date_dd + "-" + old_date_match.group(3) + "-" + old_date_match.group(5)))
            return str(old_date_dd + "-" + old_date_match.group(3) + "-" + old_date_match.group(5))
            break
    return "No date"
                                      

    #Check if a certain word occurs in the text (this version also supports word part of a larger word, e.g. "økse"-keyword matches both "økse" and "øksehoved")
def check_word_in_text(text, word):
    pattern = re.compile(rf"{re.escape(word)}", flags = re.IGNORECASE)
    return bool(pattern.search(text))
    

def extract_aut_number_from_table(doc):
    # Extract autopsy number from the first table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Extract and return the CPR number
                aut_no_match = re.search(r"04\.01\.\d{1,3}\.\d{2}", cell.text)
                #aut_no_match = re.search(r"04\.01\.d{1,3}\.\d{2}", cell.text)
                if aut_no_match:
                    print (str(aut_no_match.group(0)))
                    return "J" + str(aut_no_match.group(0))
                else:
                    OK_no_match = re.search(r"OK(\ ?)(\d{1,3}(\\|\-|\/|\ )\d{2,4}|\d{3,5})", cell.text)
                    if OK_no_match:
                        print(str(OK_no_match.group(0)))
                        return "J" + str(OK_no_match.group(0))
                        break

    for paragraph in doc.paragraphs:
        OK2_no_match = re.search(r"OK(\ ?)(\d{1,3}(\\|\-|\/|\ )\d{2,4}|\d{3,5})", paragraph.text)
        if OK2_no_match:
            print(str(OK2_no_match.group(0)))
            return "J" + str(OK2_no_match.group(0))
            break

                    
    return "No match"

#skal ændres, så kun forrådnelse i udv. US kommer med - bruge dødsstivhed?
def extract_putrefaction(text):
    #Extract putrefaction yes/no from external exam
    putrefaction_pattern = re.compile(r"(\w+ \w+ \w+) forrådnelse", re.IGNORECASE)

    putrefaction = None

    sentences = re.split(r"[.!?]", text)

    #find indices of the phrases between "udvendig undersøgelse" and "indvendig undersøgelse"
    start_index = next((i for i, sentence in enumerate(sentences) if re.search(r"udvendig undersøgelse", sentence, re.IGNORECASE)), None)
    end_index = next((i for i, sentence in enumerate(sentences) if re.search(r"indvendig undersøgelse", sentence, re.IGNORECASE)), None)

    # Subset sentences between the two phrases
    if start_index is not None and end_index is not None:
        subset_sentences = sentences[start_index + 1:end_index]

    # Find the index of "foreligger"
    foreligger_index = next((i for i, sentence in enumerate(sentences) if re.search(r"foreligger", sentence, re.IGNORECASE)), None)

    # Include sentences from the beginning until "foreligger"
    try:
        if foreligger_index is not None:
            beginning_sentences = sentences[:foreligger_index + 1]
            combined_sentences = beginning_sentences + subset_sentences
    except Exception as e:
        print(f"Error processing foreligger")

    
    # Include final 4 sentences in document
    try:
        end_sentences = sentences[-4:]
        final_sentences = combined_sentences + end_sentences
    except Exception as e:
        print(f"Error processing end sentences")
        

    try:
        putre_sentences = [s.strip() for s in final_sentences if "forrådnelse" in s.lower() or "grønlig misfarvning" in s.lower()]
        putrefaction = putre_sentences
    except Exception as e:
        print(f"Error processing putre_sentences")

    return putrefaction

# Extract putrefaction degree using standard phrases
def putrefaction_degree(text):
    priority = {
        "PRONOUNCED": 1,
        "MODERATE": 2,
        "DISCREET": 3,
        "NONE": 4,
        "INSECT ONLY": 5,
        "NO MATCH": 6
    }

    sentences = re.split(r"[.!?]", text)

    start_index = next((i for i, sentence in enumerate(sentences) if re.search(r"udvendig undersøgelse", sentence, re.IGNORECASE)), None)
    end_index = next((i for i, sentence in enumerate(sentences) if re.search(r"indvendig undersøgelse", sentence, re.IGNORECASE)), None)

    subset_sentences = sentences[start_index + 1:end_index] if start_index is not None and end_index is not None else []

    foreligger_index = next((i for i, sentence in enumerate(sentences) if re.search(r"foreligger", sentence, re.IGNORECASE)), None)
    beginning_sentences = sentences[:foreligger_index + 1] if foreligger_index is not None else []

    combined_sentences = beginning_sentences + subset_sentences
    final_sentences = combined_sentences + sentences[-4:]

    putre_sentences = [s.strip() for s in final_sentences if "forrådnelse" in s.lower() or "grønlig" in s.lower()]

    if not putre_sentences:
        return "NO MENTION"

    found_levels = []

    for sentence in putre_sentences:
        if re.search(r"\bsvær?e\b", sentence, re.IGNORECASE) or \
           re.search(r"\budtalt(e)?\b", sentence, re.IGNORECASE) or \
           re.search(r"\bfremskreden\b", sentence, re.IGNORECASE):
            found_levels.append("PRONOUNCED")
        elif re.search(r"\bmoderat\b", sentence, re.IGNORECASE) or \
             re.search(r"\bmoderate\b", sentence, re.IGNORECASE) or \
             re.search(r"\bmiddelsvær\b", sentence, re.IGNORECASE) or \
             re.search(r"\b(hud)?afløsning\b", sentence, re.IGNORECASE):
            found_levels.append("MODERATE")
        elif re.search(r"\bgrøn(lig )?(mis)?farvning\b", sentence, re.IGNORECASE) or \
             re.search(r"\blet?t?e\b", sentence, re.IGNORECASE) or \
             re.search(r"\bkartegning\b", sentence, re.IGNORECASE) or \
             re.search(r"\bbegyndende\b", sentence, re.IGNORECASE):
            found_levels.append("DISCREET")
        elif re.search(r"\bingen\b", sentence, re.IGNORECASE) or \
             re.search(r"\bikke\b", sentence, re.IGNORECASE):
            found_levels.append("NONE")
        elif re.search(r"insektangreb|maddike", sentence, re.IGNORECASE):
            found_levels.append("INSECT ONLY")

    if not found_levels:
        return "NO MATCH"

    # Return the most severe level based on priority
    return sorted(found_levels, key=lambda x: priority[x])[0]


def extract_age(doc):
    #Extract age from conclusion
    stillborn_pattern  = re.compile(r"\bdødfødt(e)?")
    newborn_pattern = re.compile(r"\bnyfødt(e)?")
    age_pattern = re.compile(r"\b(\d{1,3})(?=-årig(e)?\b)", re.IGNORECASE)
    age_month_pattern = re.compile(r"\b(\d{1,3})(?= måneder gam(mel|le)?\b)", re.IGNORECASE)
    age_week_pattern = re.compile(r"\b(\d{1,2})(?= uger gam(mel|le)?\b)", re.IGNORECASE)
    age_day_pattern = re.compile(r"\b(\d{1,2})(?= dage gam(mel|le)?\b)", re.IGNORECASE)
    fetal_week_pattern = re.compile(r"\b(fosteruge) (\d{1,2})", re.IGNORECASE)

    age = None
    age_unit = None

    
    stillborn_match = stillborn_pattern.search(doc)
    if stillborn_match:
        age = "."
        age_unit = "stillborn"
        return age, age_unit

    newborn_match = newborn_pattern.search(doc)
    if newborn_match:
        age = "."
        age_unit = "newborn"
        return age, age_unit

    fetal_week_match = fetal_week_pattern.search(doc)
    if fetal_week_match:
        age = fetal_week_match.group(2)
        age_unit = "fetal weeks"
        
    age_month_match = age_month_pattern.search(doc)
    if age_month_match:
        age = age_month_match.group(0)
        age_unit = "mon"
        return age, age_unit

    age_week_match = age_week_pattern.search(doc)
    if age_week_match:
        age = age_week_match.group(0)
        age_unit = "wk"
        return age, age_unit

    age_day_match = age_day_pattern.search(doc)
    if age_day_match:
        age = age_day_match.group(0)
        age_unit = "days"
        return age, age_unit
    
    age_match = age_pattern.search(doc)
    if age_match:
        age = age_match.group(0)
        age_unit = "yrs"
        return age, age_unit

        
    return age, age_unit

def extract_sex(text):
    #Extract gender from conclusion
    sex_pattern = re.compile(r"(årig[^.]*(mand|kvinde))|((gammel|gamle|årige)[^.]*(dreng|pige))")

    sex = None
    sexAdult = None
    sexChild = None

    sex_match = sex_pattern.search(text)
    if sex_match:
        sexAdult = sex_match.group(2)
        sexChild = sex_match.group(5)
    if sexAdult == "mand":
        sex = "M"
    if sexAdult == "kvinde":
        sex = "K"
    if sexChild == "dreng":
        sex = "M"
    if sexChild == "pige":
        sex = "K"

    return sex

def extract_supp(text):
    #Determine if record i primary or supplementary report
    supp_pattern = re.compile(r"\bsupplerende erklæring til(?: retslægelig)? obduktion|obduktion-supl\b", re.IGNORECASE)

    supp = None

    supp_match = supp_pattern.search(text)
    if supp_match:
        supp = "Supp"
    else:
        supp == "Prim"

    return supp

def read_list_from_file(filename):
    #Reads a list of items from a text file, one item per line
    with open(filename, "r", encoding="utf-8") as file:
        return [line.strip() for line in file if line.strip()]

def extract_lesions(doc):
    para_v = None
    para_s = None
    para_l = []
    list_les = read_list_from_file(r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\lesion_lists\list_les.txt")
    list_col = read_list_from_file(r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\lesion_lists\list_col.txt")
    list_loc = read_list_from_file(r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\lesion_lists\list_loc.txt")
    list_sha = read_list_from_file(r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\lesion_lists\list_sha.txt")


    for i, para in enumerate(doc.paragraphs[30:], start = 30):
        if para_v is None and "tegn på vold" in para.text:
            para_v = i
        elif para_v is not None and "Indvendig" in para.text:
            para_s = i
            break

    if para_v is not None and para_s is not None and para_v < para_s:
        para_l = [para.text for para in doc.paragraphs[para_v + 1:para_s]]

    lesion_dict = {}
    lesion_count = 0

    #Iterate trough each paragraph in the list.
    for para in para_l:
        #Make a list to store found lesions in current paragraph
        found_lesions = []
        for lesion in list_les:
            if lesion in para:
                found_lesions.append(lesion)
        if found_lesions:
            lesion_count += 1
            #Concatenate found lesions with " / " and store in dictionary
            lesion_dict[f"lesion_{lesion_count}"] = " / ".join(found_lesions)

            #Search for location words in the same paragraph using regex
            found_locs = []
            for loc in list_loc:
                #Regex pattern to match color word AND inflections (e.g. the word with an extra letter, e.g. skulder / skulders
                pattern = re.compile(r"\b" + re.escape(loc[:-1]) + r"\w*\b", re.IGNORECASE)
                if pattern.search(para):
                    found_locs.append(loc)
            if found_locs:
                lesion_dict[f"lesion_{lesion_count}_loc"] = " / ".join(found_locs)

            #Search for color words in the same paragraph using regex
            found_colors = []
            for color in list_col:
                #Regex pattern to match color word AND inflections (e.g. the word with an extra letter, e.g. rødlig / rødlige / rødligt
                pattern = re.compile(r"\b" + re.escape(color[:-1]) + r"\w*\b", re.IGNORECASE)
                if pattern.search(para):
                    found_colors.append(color)
            if found_colors:
                lesion_dict[f"lesion_{lesion_count}_col"] = " / ".join(found_colors)

            #Search for shape words in the same paragraph using regex
            found_shapes = []
            for shape in list_sha:
                #Regex pattern to match shape word AND inflections (e.g. the word with an extra letter, e.g. afrundet / afrundede
                pattern = re.compile(r"\b" + re.escape(shape[:-1]) + r"\w*\b", re.IGNORECASE)
                if pattern.search(para):
                    found_shapes.append(shape)
            if found_shapes:
                lesion_dict[f"lesion_{lesion_count}_sha"] = " / ".join(found_shapes)

            #Search for dimensions after the lesion words
            found_dimensions = []
            for lesion in found_lesions:
                lesion_index = para.find(lesion)
                if lesion_index != -1:
                    #Extract text after the lesion word
                    text_after_lesion = para[lesion_index + len(lesion):]
                    #Find all dimension patterns
                    pattern = re.compile(r"\b\d+(?:,\d+)?(?:\s*x\s*\d+(?:,\d+)?)?\b")
                    for match in pattern.finditer(text_after_lesion):
                        found_dimensions.append(match.group())
                if found_dimensions:
                    found_dimensions = list(dict.fromkeys(found_dimensions))
                    lesion_dict[f"lesion_{lesion_count}_dim"] = " / ".join(found_dimensions)
                
    lesion_dict["lesion_count"] = lesion_count

    return lesion_dict

def process_documents(folder_path, keywords, keyword_COD, keyword_2_COD):
    # Initialize a list to store dictionaries of data for each document
    print("Processsing docx-documents!")
    all_data = []
    all_keys = OrderedDict()
    num_files_processed = 0
    total_time = 0

    # List to store all paths for docx-files
    docx_files = []

    # List to store all paths for pdf-files
    pdf_files = []

    # Loop through all files in the specified folder and subfolders (using os.walk)
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                docx_files.append(os.path.join(root, file))
                start_time = time.time()
            #elif file.endswith(".pdf"):
            #    pdf_files.append(os.path.join(root, file))
            #    start_time = time.time()

    # Total number of docx-files an pdf-files
    total_files = len(docx_files) + len(pdf_files)
    print("The total number of files is: " + str(total_files))
    print("The total number of word-files is: " + str(len(docx_files)))
    print("The total number of pdf-files is: " + str(len(pdf_files)))
    time.sleep(2)
    
    # Loop through all docx-files
    for file_path in docx_files:
        filename = os.path.basename(file_path)

        # Read the document using the docx module
        try:
            start_time = time.time()
            doc = Document(file_path)
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            continue

        # Concatenate all paragraph texts into a single string
        doc_text = " ".join([paragraph.text.replace("\n", " ").replace("\r", " ") for paragraph in doc.paragraphs])
        print(len(doc_text))

        # Extract CPR number from table in the document
        cpr_number = extract_cpr_number_from_table(doc)

        # Extract autopsy date from table in the document
        aut_date = extract_aut_date_from_table(doc)
            
        # Extract autopsy record number from table in the document
        aut_number = extract_aut_number_from_table(doc)

        # Extract if primary or supplementary report from the concatenated text
        supp = extract_supp(doc_text)
        
        # Extract lung weights from the concatenated text
        weights = extract_lung_weights(doc_text, keywords)

        # Extract organ sizes
        organ_sizes = extract_organ_size(doc_text, organ_keywords)

        # Extract wall thicknesses from the concatenated text
        thicknesses = extract_wall_thicknesses(
            doc_text, ["højre hjertekammer", "venstre hjertekammer", "hjerteskille"]
        )

        # Extract pleural fluid volume from the concatenated text
        volumes = extract_pleural_fluid(doc_text, ["højre","venstre","bughule"])

        # Extract height and weight from the concatenated text
        height, bod_weight, bod_weight_unit = extract_height_weight(doc_text)

        # Extract putrefaction from the concatenated text
        putrefaction = extract_putrefaction(doc_text)

        # Extract putrefaction level from text
        putre_level = putrefaction_degree(doc_text)

        # Extract keyword from text
        keyword = check_word_in_text(doc_text, "autoerot")

        # Extract age from the concatenated text
        age, age_unit = extract_age(doc_text)

        # Extract sex from the concatenated text
        sex = extract_sex(doc_text)

        # Check if COD keywords in the given list is present in the document, create dictionary
        regex_dict = {
            "uoplyst": r"ikke oplyst|uoplyst",
            "drukning": r"drukning",
            "hjertesvigt": r"akut hjertesvigt",
            "forgiftning": r"forgiftning(?![^.]+(kulilte|cyanid))",
            "hængning": r"hængning",
            "skud": r"skudlæsion",
            "stik_snit": r"stiklæsion|snitlæsion|stiksår|snitsår",
            "forblødning": r"forblødning",
            "forbrænding": r"forbrænding",
            "lungebetændelse": r"lunge[^.]+betændelse|betændelse[^.]+lunge|lungebetændelse",
            "ikke holdepunkt": r"ikke holdepunkt",
            "supp_no_change": r"resultat[^.]+giver ikke|resultat[^.]+ændrer ikke",
        }
        keyword_COD_dict = search_for_COD_keywords(doc_text, regex_dict)
        
        # Look up COD paragraph and store whole paragraph as text variable
        textCOD_dict = store_COD_text(doc_text)

        # Look up vaccination sentences and store each sentence as text variable - up to two sentences
        textVAC = store_vaccine_text(doc).replace("\n", " ")

        # Look up findesteds paragraph and store whole paragraph as text variable
        finde_text = store_finde_text(doc)

        # Findeomstændigheder - dictionary of terms and associated regexes, that are checked in paragraphs with phrases such as "af disse papirer", "nu afdøde", etc. - check the function for all terms
        findeomst_dict = {
            "fundet_i_vand": (
                r"(fundet|livløs|\bfandt|liggende|ligget|lå|nedsunket|under|\bflydende\b|drivende|bunden|ude i vandet|fik i|trukket op|optaget i|spottet|bjerget|fisket op|reddet)"
                r"(?:(?![.,:]\s).)+(\bflydende|druknet|drivende|i vandet|af vandet|i en å|\bå\b|\bbrønd\b|sivbrønd|sivområde|trawl|under vand|lavt vand|swimmingpool|vandkanten"
                r"|vandoverfladen|vandhul|fra båden|bælt\b|vandløb|dam\b|fiskedam|\bsø\b|søen\b|gadekær|på bunden|havbunden|saltvandsbassin|havnebassin|drevet i land"
                r"|vandhul|bundgarn|farvand|fjord|voldgrav|strandkanten|havet\b|havstokken|\bkanal|\bhavn(?!et))"
            ),
            "trafik": r"påkørt|fører af|passager\b|færdselsuheld|trafikuheld|trafikulykke"
        }

        findeomst_result = findeomst(doc, findeomst_dict)

        # Look for keyCT in paragraphs with "CT" and following four paragraphs
        keyCT_present = CT_search(doc, keywordCT)

        # Look for "skumsvamp" in paragraphs and return all paragraphs where this is true
        skum_para = skumsvampPara(doc)

        # Look for "strip" in the whole document, return all paragraphs where "strip" is found as a single string
        strip_text = stripPara(doc)

        # Look for "tegn på sygdom*" in paragraphs and return all text in paragraph after the phrase
        TPS = search_TPS(doc).replace("\n", " ")

        #Look for "efter det oplyste" and return subsequent text until finding "mand|kvinde"
        KS = kendtMed(doc)

        #Look for "hjerteposen" and return all text in that paragraph
        textHeart = hjerteText(doc).replace("\n", " ")

        #Look for "Legemspulsåren og" and return all text in that paragraph
        textAorta = aortaText(doc).replace("\n", " ")

        #Look for "Halspulsårerne afgår" and return all text in that paragraph
        textCarotid = carotidText(doc).replace("\n", " ")

        #Compile list of paragraphs with lesion data
        #lesions = extract_lesions(doc)

        # Create a dictionary to store the data for this document
        data = {
            "File Name": filename,
            "CPR Number": cpr_number,
            "aut_number": aut_number,
            "Prim_status": supp,
            "Autopsy Date": aut_date,
            "Age": age,
            "Age unit": age_unit,
            "Sex": sex,
            **weights,  # Unpack the lung weights dictionary
            **organ_sizes, #Unpack the organ sizes dictionary
            **thicknesses,  # Unpack the wall thicknesses dictionary
            **volumes, # Unpack the pleural fluid volumes dictionary
            "Højde": height,
            "Vægt": bod_weight,
            "Vægtenhed": bod_weight_unit,
            "Putrefaction": putrefaction,
            "Putre_level": putre_level,
            "Autoerot": keyword,
            **keyword_COD_dict,  # Unpack COD keyword matches
            "COD tekst": textCOD_dict,
            "Finde tekst": finde_text,
            "Vaccine text": textVAC,
            **findeomst_result,
            "keywordCT: "+str(keywordCT): keyCT_present,
            "Skumsvamp tekst": skum_para,
            "Strip_text": strip_text,
            "TPS": TPS,
            "Kendte sygdomme": KS,
            "Hjertebeskrivelse": textHeart,
            "Aortabeskrivelse": textAorta,
            "Carotider_beskrivelse": textCarotid,
            #"LW/HW": (weights.get("venstre lunge") + weights.get("Højre lunge"))/weights.get("Hjertet"),
            #**lesions #unpack the lesions dictionary
        }

        for key in data.keys():
            if key not in all_keys:
                all_keys[key] = None

        # Append the data dictionary to the list
        all_data.append(data)

        # Calculate processing time for processed file
        elapsed_time = time.time() - start_time
        total_time += elapsed_time
        num_files_processed += 1

        # Compute the running average of file processing time
        average_time_per_file = total_time / num_files_processed

        # Estimate the remaining time
        remaining_files = total_files - num_files_processed
        remaining_time = average_time_per_file * remaining_files
        remaining_time_hms = time.strftime("%H:%M:%S", time.gmtime(remaining_time))

        # Calculate progress percentage
        progress_percentage = (num_files_processed / len(docx_files)) * 100

        # Print the progress percentage and estimated remaining time
        print(f"Conversion {progress_percentage:.2f}% complete.")
        print(f"Expected time left to completion: {remaining_time_hms}")
        
        print(filename)

    # Sort lesion columns for each document
    #for data in all_data:
        #Extract the "lesion_" keys and sort them
        #lesion_keys = sorted(
        #    (key for key in data.keys() if key.startswith("lesion_") and key != "lesion_count"),
        #    key=lambda x: int(x.split("_")[1])  #Extract the lesion number for sorting
        #)

        #Rebuild the dictionary with sorted lesion columns
        #sorted_data = {"lesion_count": data.pop("lesion_count")}
        #for key in lesion_keys:
        #    sorted_data[key] = data[key]
        #Include other fields not related to lesions
        #sorted_data.update({key: value for key, value in data.items() if not key.startswith("lesion_")})

        #Replace the original entry with the sorted data
        #data.clear()
        #data.update(sorted_data)

    return all_data, list(all_keys.keys())
    

# export_to_csv now uses the "all_keys" variable to create the field names, so even if first document is missing values, it should not produce an error
# ADDED 2025-08-08 export_to_csv now finds duplicate CPR numbers. If the have the same aut_number, only the one with highest "File Name" is kept. If there are multiple aut_num, all duplicates are kept.
# A log file with duplicates, including which are removed, are created and stored in a separate CSV-file. 
def export_to_csv(data, all_keys, csv_filename):
    # Group entries by CPR Number
    cpr_groups = {}
    for entry in data:
        cpr_number = entry.get("CPR Number")
        if cpr_number not in cpr_groups:
            cpr_groups[cpr_number] = []
        cpr_groups[cpr_number].append(entry)

    filtered_data = []
    duplicates_log = []

    for cpr_number, entries in cpr_groups.items():
        if len(entries) == 1:
            # Only one entry, keep it — no logging
            filtered_data.append(entries[0])
            continue

        # Group by aut_number
        aut_groups = {}
        for entry in entries:
            aut_number = entry.get("aut_number")
            if aut_number not in aut_groups:
                aut_groups[aut_number] = []
            aut_groups[aut_number].append(entry)

        for aut_number, aut_entries in aut_groups.items():
            if len(aut_entries) == 1:
                # Unique aut_number within duplicate CPR group — keep it
                filtered_data.append(aut_entries[0])
                duplicates_log.append({
                    "File Name": aut_entries[0].get("File Name"),
                    "CPR Number": cpr_number,
                    "aut_number": aut_number,
                    "Omitted": "No"
                })
            else:
                # Multiple entries with same CPR and aut_number
                aut_entries_sorted = sorted(aut_entries, key=lambda x: x.get("File Name"))
                for entry in aut_entries_sorted[:-1]:
                    duplicates_log.append({
                        "File Name": entry.get("File Name"),
                        "CPR Number": cpr_number,
                        "aut_number": aut_number,
                        "Omitted": "Yes"
                    })
                # Keep the last one (highest File Name)
                filtered_data.append(aut_entries_sorted[-1])
                duplicates_log.append({
                    "File Name": aut_entries_sorted[-1].get("File Name"),
                    "CPR Number": cpr_number,
                    "aut_number": aut_number,
                    "Omitted": "No"
                })

    # Write the filtered data to the main CSV file
    with open(csv_filename, "w", newline="", encoding="utf-16") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=all_keys, quoting=csv.QUOTE_ALL)
        writer.writeheader()
        writer.writerows(filtered_data)

    # Write the duplicates log to a separate CSV file
    if duplicates_log:
        with open("duplicates.csv", "w", newline="", encoding="utf-8") as log_file:
            log_writer = csv.DictWriter(log_file, fieldnames=["File Name", "CPR Number", "aut_number", "Omitted"])
            log_writer.writeheader()
            log_writer.writerows(duplicates_log)
        
# Example usage:
folder_path = r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\Primære erklæringer 1992-2024"
#folder_path = r"S:\RPA\7. Retspatologi\Andet\JOB_automatiskdataudtræk_CJW\Workspace\test"
keywords = ["Højre lunge", "Venstre lunge", "Hjerte", "Milt", "Leveren", "Hjernen", "Højre nyre", "Venstre nyre"]
organ_keywords = ["Hjertet", "Leveren", "Højre nyre", "Venstre nyre"]
output_csv_filename = "output_2025_08_26_supp.csv"
keyword_COD = "drukning"  # Define keywordCOD
keyword_2_COD = "akut hjertesvigt"
keyword_3_COD = "forgiftning"
keywordCT = "hjertepose" #Define keyword to look for in CT-paragraphs
result, keys = process_documents(folder_path, keywords, keyword_COD, keyword_2_COD)

# Export the result to a CSV file
export_to_csv(result, keys, output_csv_filename)
# Write your code here :-)

